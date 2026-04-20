"""Build CFD Strategy Word template (blank + filled sample). Rey Capital themed."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn, nsmap
from docx.oxml import OxmlElement

OUT = r"C:/Users/NinadKulkarni/PycharmProjects/AlgoStrategies/docs/strategy-intake"
LOGO = f"{OUT}/ReyCapital_Logo.png"

# Rey Capital palette (sampled from logo)
NAVY = RGBColor(0x10, 0x26, 0x7A)     # deep navy banner
BLUE = RGBColor(0x1E, 0x3F, 0xAE)     # primary Rey brand blue
ACCENT = RGBColor(0x34, 0x57, 0xC7)   # lighter accent
GREY_TXT = RGBColor(0x40, 0x40, 0x40)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
CONF_RED = RGBColor(0xC0, 0x00, 0x00)

FONT = "Calibri"


# ---------------------- helpers ----------------------

def set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_run(run, size=11, bold=False, color=None, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = color


def banner(doc, text, color=NAVY):
    """Full-width colored banner with white text (single-cell table)."""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.autofit = False
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "{:02X}{:02X}{:02X}".format(*color))
    cell.width = Cm(17)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    set_run(r, size=20, bold=True, color=WHITE)
    # padding
    for side in ("top", "bottom"):
        p_pr = p._p.get_or_add_pPr()
        sp = OxmlElement("w:spacing")
        sp.set(qn("w:before"), "120")
        sp.set(qn("w:after"), "120")
        p_pr.append(sp)
        break


def h1(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)
    # bottom border
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:color"), "1E3FAE")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)
    r = p.add_run(text)
    set_run(r, size=16, bold=True, color=BLUE)
    return p


def h2(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run(r, size=13, bold=True, color=ACCENT)
    return p


def para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_run(r, size=size, bold=bold, color=GREY_TXT)
    r.italic = italic
    return p


def bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    r = p.runs[0] if p.runs else p.add_run(text)
    if not p.runs:
        r = p.add_run(text)
    else:
        r.text = text
    set_run(r, size=11, color=GREY_TXT)
    return p


def field_table(doc, rows, col1_width=Cm(6), col2_width=Cm(11), placeholder="__________________________"):
    """Two-column 'label | value' table used across the doc."""
    tbl = doc.add_table(rows=len(rows), cols=2)
    tbl.autofit = False
    for i, (label, value) in enumerate(rows):
        lcell = tbl.rows[i].cells[0]
        vcell = tbl.rows[i].cells[1]
        lcell.width = col1_width
        vcell.width = col2_width
        set_cell_shading(lcell, "1E3FAE")
        set_cell_shading(vcell, "E4EAFB" if i % 2 == 0 else "FFFFFF")
        # label
        p = lcell.paragraphs[0]
        r = p.add_run(label)
        set_run(r, size=10.5, bold=True, color=WHITE)
        # value
        p2 = vcell.paragraphs[0]
        r2 = p2.add_run(value if value else placeholder)
        set_run(r2, size=10.5, bold=False, color=GREY_TXT)
        # vertical align
        lcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        vcell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _apply_borders(tbl)
    return tbl


def grid_table(doc, headers, rows, col_widths=None):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.autofit = False
    # header
    for i, hdr in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_shading(cell, "1E3FAE")
        if col_widths:
            cell.width = col_widths[i]
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(hdr)
        set_run(r, size=10.5, bold=True, color=WHITE)
    # rows
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            if col_widths:
                cell.width = col_widths[ci]
            set_cell_shading(cell, "E4EAFB" if ri % 2 else "FFFFFF")
            p = cell.paragraphs[0]
            r = p.add_run(str(val))
            set_run(r, size=10, color=GREY_TXT)
    _apply_borders(tbl)
    return tbl


def _apply_borders(tbl):
    tbl_pr = tbl._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "B4B4B4")
        borders.append(b)
    tbl_pr.append(borders)


def page_break(doc):
    doc.add_page_break()


def set_defaults(doc):
    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(11)
    style.font.color.rgb = GREY_TXT
    # page margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)


def footer(doc, text):
    for section in doc.sections:
        f = section.footer.paragraphs[0]
        f.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = f.add_run(text)
        set_run(r, size=9, color=GREY_TXT)
        r.italic = True


def apply_header(doc):
    """Add CONFIDENTIAL stripe + Rey Capital tag in every page header."""
    for section in doc.sections:
        hdr = section.header.paragraphs[0]
        hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # left tag
        r1 = hdr.add_run("REY CAPITAL  |  ")
        set_run(r1, size=9, bold=True, color=BLUE)
        r2 = hdr.add_run("CONFIDENTIAL — INTERNAL USE ONLY")
        set_run(r2, size=9, bold=True, color=CONF_RED)
        r2.italic = True
        # bottom border on header
        p_pr = hdr._p.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "6")
        bottom.set(qn("w:color"), "1E3FAE")
        p_bdr.append(bottom)
        p_pr.append(p_bdr)


def insert_logo(doc, width_cm=6.5):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    r.add_picture(LOGO, width=Cm(width_cm))
    return p


def confidential_banner(doc):
    tbl = doc.add_table(rows=1, cols=1)
    cell = tbl.cell(0, 0)
    set_cell_shading(cell, "FDE7E9")
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("CONFIDENTIAL — REY CAPITAL INTERNAL USE ONLY  |  DO NOT DISTRIBUTE")
    set_run(r, size=11, bold=True, color=CONF_RED)
    r.italic = True
    _apply_borders(tbl)


# ==================================================================
# CONTENT BUILDER
# ==================================================================

PH = "__________________________"  # placeholder for blank template


def build(path, sample=False):
    doc = Document()
    set_defaults(doc)
    apply_header(doc)

    # ---------------- COVER ----------------
    insert_logo(doc, width_cm=7.5)
    para(doc, "")
    confidential_banner(doc)
    para(doc, "")
    banner(doc, "CFD STRATEGY SPECIFICATION DOCUMENT")
    para(doc, "Rey Capital | Smart Investments — Manual Trader → Automation Intake", italic=True, size=11)
    para(doc, "")
    banner(doc, "STRATEGY METADATA", color=BLUE)

    meta = [
        ("Strategy Name",        "EMA 9-21 Crossover (FX & Indices)"     if sample else PH),
        ("Strategy Code",        "CFD_EMA_9_21_v1"                        if sample else PH),
        ("Version",              "1.0"                                    if sample else PH),
        ("Author / PM",          "Ninad Kulkarni"                         if sample else PH),
        ("Desk",                 "Systematic FX Desk"                     if sample else PH),
        ("Submission Date",      "2026-04-20"                             if sample else PH),
        ("Strategy Category",    "Trend-Following"                        if sample else PH),
        ("Direction",            "Both (Long & Short)"                    if sample else PH),
        ("Base Currency",        "USD"                                    if sample else PH),
        ("Primary Timeframe",    "H1"                                     if sample else PH),
        ("Companion Excel File", "CFD_Strategy_Params_SAMPLE_EMA_9_21_Crossover.xlsx" if sample else "CFD_Strategy_Params_v1.xlsx"),
        ("Document Status",      "DRAFT — Pending Sign-off"               if sample else "BLANK TEMPLATE"),
    ]
    field_table(doc, meta)
    para(doc, "")
    para(doc,
         "Fill every field. Ambiguous, vague, or missing sections will cause the automation team to reject this "
         "specification. If a section does not apply, write 'N/A' with a brief justification. "
         "This document is the contract between the trader and the automation team — rules stated here are what will "
         "be coded. Anything not stated here will NOT be implemented.",
         italic=True, size=10)
    page_break(doc)

    # ---------------- 1. IDENTIFICATION ----------------
    h1(doc, "1. Strategy Identification")
    h2(doc, "1.1 Core Thesis")
    if sample:
        para(doc,
             "Markets exhibit persistent short-to-medium term momentum on liquid FX majors and equity indices. "
             "When the 9-period EMA crosses the 21-period EMA on the H1 timeframe, it signals a shift in short-term "
             "momentum that, when filtered by session and volatility regime, captures 1R–3R directional moves before "
             "mean-reversion sets in. Edge is strongest during London and London-NY overlap sessions when institutional "
             "flow dominates.")
    else:
        para(doc, "[ One-line thesis — why does this strategy make money? What is the edge hypothesis? ]", italic=True)
        for _ in range(3):
            para(doc, PH)

    h2(doc, "1.2 Edge Source & Academic Support")
    if sample:
        bullet(doc, "Momentum anomaly (Jegadeesh & Titman, 1993) — cross-asset.")
        bullet(doc, "Session-based volatility clustering — London/NY overlap provides highest signal-to-noise.")
        bullet(doc, "Moving-average crossover as trend filter — reduces false signals in ranging markets.")
    else:
        para(doc, "[ Cite papers, prior research, or internal evidence of edge. ]", italic=True)
        for _ in range(3):
            bullet(doc, PH)

    # ---------------- 2. MARKET & INSTRUMENT SCOPE ----------------
    h1(doc, "2. Market & Instrument Scope")
    h2(doc, "2.1 Asset Classes")
    if sample:
        bullet(doc, "FX Majors (high liquidity, tight spreads)")
        bullet(doc, "Commodity CFDs (Gold only)")
        bullet(doc, "Index CFDs (US and EU cash indices)")
    else:
        for _ in range(3):
            bullet(doc, PH)

    h2(doc, "2.2 Instrument Universe")
    if sample:
        grid_table(doc,
                   ["Symbol", "Asset Class", "Typical Spread", "Why Included"],
                   [
                       ["EURUSD", "FX Major",     "0.8 pips", "Highest FX liquidity, lowest cost"],
                       ["GBPUSD", "FX Major",     "1.2 pips", "Strong trends during London"],
                       ["USDJPY", "FX Major",     "1.0 pips", "Carry & BoJ-driven trends"],
                       ["XAUUSD", "Commodity",    "3.0 pips", "Persistent macro trends"],
                       ["US30",   "Index CFD",    "2.5 pts",  "US equity momentum proxy"],
                       ["GER40",  "Index CFD",    "1.8 pts",  "EU equity momentum proxy"],
                   ],
                   col_widths=[Cm(2.5), Cm(3.5), Cm(3), Cm(7)])
        para(doc, "Full per-symbol specs: see Params.xlsx → Instruments tab.", italic=True, size=10)
    else:
        grid_table(doc,
                   ["Symbol", "Asset Class", "Typical Spread", "Why Included"],
                   [[PH]*4 for _ in range(6)],
                   col_widths=[Cm(2.5), Cm(3.5), Cm(3), Cm(7)])

    h2(doc, "2.3 Liquidity Filters")
    if sample:
        bullet(doc, "Minimum 10 billion USD average daily volume (FX).")
        bullet(doc, "Spread must be within 1.5x of typical spread at time of signal.")
        bullet(doc, "No trading in first/last 15 minutes of any session.")
    else:
        for _ in range(3):
            bullet(doc, PH)

    # ---------------- 3. TIMEFRAME & SESSION ----------------
    h1(doc, "3. Timeframe & Session")
    rows = [
        ("Primary Timeframe",      "H1"                       if sample else PH),
        ("Confirmation Timeframe", "H4 (trend alignment)"     if sample else PH),
        ("Session Window (GMT)",   "07:00 – 20:00"            if sample else PH),
        ("Preferred Sessions",     "London, London-NY Overlap" if sample else PH),
        ("No-Entry Window",        "14:25–14:35 GMT (US data releases)" if sample else PH),
        ("Holding Period",         "Intraday to 3 days max"   if sample else PH),
        ("Weekly Square-Off",      "21:00 GMT Friday"         if sample else PH),
        ("Blackout Dates",         "FOMC days, NFP, CPI, ECB" if sample else PH),
    ]
    field_table(doc, rows)

    # ---------------- 4. ENTRY LOGIC ----------------
    h1(doc, "4. Entry Logic")
    h2(doc, "4.1 Indicators Used")
    grid_table(doc,
               ["Indicator", "Parameters", "Source / Price Input"],
               ([
                   ["EMA (fast)",   "Period = 9",  "Close"],
                   ["EMA (slow)",   "Period = 21", "Close"],
                   ["ATR",          "Period = 14", "High-Low-Close"],
                   ["EMA (trend filter)", "Period = 200", "Close (on H4)"],
               ] if sample else [[PH, PH, PH] for _ in range(4)]),
               col_widths=[Cm(5), Cm(5), Cm(7)])

    h2(doc, "4.2 Long Entry Rules")
    if sample:
        bullet(doc, "Condition 1: EMA(9) crosses ABOVE EMA(21) on current H1 bar close.")
        bullet(doc, "Condition 2: Close price is ABOVE EMA(200) on H4 (trend filter).")
        bullet(doc, "Condition 3: ATR(14) > 70% of 20-day average ATR (volatility filter — skip dead markets).")
        bullet(doc, "Condition 4: Current time is within preferred session window (07:00–20:00 GMT).")
        bullet(doc, "Condition 5: Spread at signal time ≤ 1.5× typical spread.")
        bullet(doc, "Order placed at next bar open.")
    else:
        para(doc, "[ List ALL conditions that must be true simultaneously for a long entry. Be unambiguous. ]", italic=True)
        for _ in range(5):
            bullet(doc, PH)

    h2(doc, "4.3 Short Entry Rules")
    if sample:
        bullet(doc, "Condition 1: EMA(9) crosses BELOW EMA(21) on current H1 bar close.")
        bullet(doc, "Condition 2: Close price is BELOW EMA(200) on H4.")
        bullet(doc, "Conditions 3–5: Same as long (volatility, session, spread filters).")
        bullet(doc, "Order placed at next bar open.")
    else:
        for _ in range(5):
            bullet(doc, PH)

    h2(doc, "4.4 Worked Example")
    if sample:
        para(doc,
             "EURUSD H1 bar closing 2026-03-15 09:00 GMT: EMA(9)=1.0845, EMA(21)=1.0842 — bullish crossover. "
             "H4 EMA(200)=1.0801, close=1.0845 (above filter). ATR(14)=0.0032 vs 20-day avg ATR=0.0028 (1.14× — pass). "
             "Spread=0.9 pips (≤ 1.5×0.8=1.2 pips). Time=09:00 GMT (London session). ALL CONDITIONS MET. "
             "Long entry at next bar open (10:00 GMT) at 1.0847.")
    else:
        para(doc, "[ Provide one worked example using real historical bar values showing every condition check. ]", italic=True)
        for _ in range(4):
            para(doc, PH)

    # ---------------- 5. EXIT LOGIC ----------------
    h1(doc, "5. Exit Logic")
    rows = [
        ("Initial Stop-Loss",   "MAX(30 pips, 1.5 × ATR(14))"                       if sample else PH),
        ("Take-Profit",         "2R (2× initial stop distance)"                     if sample else PH),
        ("Trailing Stop",       "Activated at 1R profit; trail 1× ATR(14) from high/low" if sample else PH),
        ("Reverse Signal Exit", "Close immediately if opposite EMA crossover"       if sample else PH),
        ("Time-Based Exit",     "Close at 21:00 GMT Friday regardless of P&L"       if sample else PH),
        ("Emergency Exit",      "Close if spread widens to > 3× typical"            if sample else PH),
    ]
    field_table(doc, rows)

    # ---------------- 6. POSITION SIZING & RISK ----------------
    h1(doc, "6. Position Sizing & Risk")
    rows = [
        ("Capital Allocation",       "1% of NAV risk per trade"               if sample else PH),
        ("Sizing Formula",           "Lot Size = (Equity × 0.01) / (Stop Pips × Pip Value)" if sample else PH),
        ("Max Concurrent Positions", "4"                                      if sample else PH),
        ("Max Daily Loss",           "3% of NAV (kill-switch triggers)"       if sample else PH),
        ("Max Drawdown Tolerance",   "15% peak-to-trough"                     if sample else PH),
        ("Leverage Cap",             "30:1 (retail) / 100:1 (prof)"           if sample else PH),
        ("Correlation Cap",          "Max 2 open positions in correlated cluster (e.g., EURUSD+GBPUSD)" if sample else PH),
        ("Sector / Cluster Groups",  "FX-USD / Precious Metals / US Indices / EU Indices" if sample else PH),
    ]
    field_table(doc, rows)
    para(doc, "Full sizing matrix: see Params.xlsx → PositionSizing tab.", italic=True, size=10)

    # ---------------- 7. ORDER EXECUTION ----------------
    h1(doc, "7. Order Execution Rules")
    rows = [
        ("Order Type (Entry)",       "Market"                               if sample else PH),
        ("Order Type (Stop)",        "SL-M"                                 if sample else PH),
        ("Order Type (Target)",      "Limit"                                if sample else PH),
        ("Slippage Tolerance",       "1.0 pips FX / 2.0 pts indices"        if sample else PH),
        ("Partial Fill Handling",    "Accept any fill ≥ 50% of requested"   if sample else PH),
        ("Rejection Retry",          "3 retries at 200ms intervals"         if sample else PH),
        ("Smart Order Routing",      "Prefer lowest-spread venue; fallback to primary" if sample else PH),
    ]
    field_table(doc, rows)

    # ---------------- 8. EDGE CASES ----------------
    h1(doc, "8. Edge Cases & Exceptions")
    grid_table(doc,
               ["Scenario", "Handling Rule"],
               ([
                   ["Circuit limit / halt",       "Cancel pending; hold existing until reopen"],
                   ["Gap opening > 1%",           "Skip signal; wait for next bar"],
                   ["Missing / stale tick > 5s",  "Pause strategy; alert ops"],
                   ["Corporate action (index)",    "Flatten 1 day before; resume after"],
                   ["CFD swap charge (22:00 GMT)", "Account for in PnL; triple swap Wednesday"],
                   ["Broker connectivity loss",   "Attempt reconnect 5×; else market-close all"],
                   ["Weekend gap (Sunday open)",  "No new Friday entries after 19:00 GMT"],
                   ["High-impact news within 30min", "Block new entries; trail stop tight"],
               ] if sample else [[PH, PH] for _ in range(8)]),
               col_widths=[Cm(6), Cm(11)])

    # ---------------- 9. BACKTEST EVIDENCE ----------------
    h1(doc, "9. Backtest Evidence")
    rows = [
        ("Backtest Period",    "2020-01-01 to 2025-12-31 (6 years)" if sample else PH),
        ("Out-of-Sample",      "2024-01 to 2025-12 (2 years)"       if sample else PH),
        ("Data Source",        "Dukascopy M1 aggregated to H1"      if sample else PH),
        ("Costs Modeled",      "Spread + commission + swap + 0.5pip slippage" if sample else PH),
        ("Portfolio Sharpe",   "2.18"                               if sample else PH),
        ("Sortino",            "2.94"                               if sample else PH),
        ("Max Drawdown",       "7.1%"                               if sample else PH),
        ("CAGR",               "16.2%"                              if sample else PH),
        ("Win Rate",           "52.5%"                              if sample else PH),
        ("Profit Factor",      "1.68"                               if sample else PH),
        ("Total Trades",       "1,959"                              if sample else PH),
    ]
    field_table(doc, rows)
    para(doc, "Per-symbol metrics & trade log: see Params.xlsx → BacktestMetrics tab.", italic=True, size=10)

    # ---------------- 10. GO-LIVE ----------------
    h1(doc, "10. Go-Live & Monitoring")
    h2(doc, "10.1 Ramp-Up Schedule")
    grid_table(doc,
               ["Phase", "Duration", "Capital", "Gate to Next Phase"],
               ([
                   ["Paper Trading",  "6 weeks", "0% (simulated)", "Live slippage & fill rate tracking backtest ±15%"],
                   ["Pilot Live",     "4 weeks", "10% of target",  "Live Sharpe ≥ 1.0 and DD < 5%"],
                   ["Scale-Up",       "4 weeks", "50% of target",  "Live Sharpe ≥ backtest × 0.7"],
                   ["Full Allocation","Ongoing", "100%",           "Quarterly review"],
               ] if sample else [[PH]*4 for _ in range(4)]),
               col_widths=[Cm(4), Cm(3), Cm(3), Cm(7)])

    h2(doc, "10.2 Daily Monitoring KPIs")
    if sample:
        bullet(doc, "Realized vs expected PnL deviation (threshold: ±2σ over 20 trades).")
        bullet(doc, "Slippage vs backtest assumption (threshold: +50%).")
        bullet(doc, "Fill rate (threshold: < 90% triggers review).")
        bullet(doc, "Average holding time vs backtest (threshold: ±30%).")
    else:
        for _ in range(4):
            bullet(doc, PH)

    h2(doc, "10.3 Kill-Switch Triggers")
    if sample:
        bullet(doc, "Daily loss > 3% of allocation → halt for the day.")
        bullet(doc, "Weekly loss > 7% → halt pending PM + risk review.")
        bullet(doc, "5 consecutive stop-outs → halt & diagnose.")
        bullet(doc, "Broker latency > 500ms sustained > 2 min → halt.")
    else:
        for _ in range(4):
            bullet(doc, PH)

    h2(doc, "10.4 Alerting")
    if sample:
        bullet(doc, "All kill-switch events: page PM + Risk + Ops (Slack + SMS).")
        bullet(doc, "Daily PnL summary emailed 22:30 GMT.")
        bullet(doc, "Weekly performance review meeting, Mondays 09:30 IST.")
    else:
        for _ in range(3):
            bullet(doc, PH)

    # ---------------- 11. SIGN-OFFS ----------------
    h1(doc, "11. Sign-offs")
    para(doc, "This strategy may not proceed to paper trading until all signatures are captured.", italic=True, size=10)
    grid_table(doc,
               ["Role", "Name", "Signature / Date", "Status"],
               ([
                   ["Trader / PM",           "Ninad Kulkarni",      "_______________ / 2026-04-20", "Pending"],
                   ["Automation Lead",       PH,                    "_______________ / __________", "Pending"],
                   ["Risk Manager",          PH,                    "_______________ / __________", "Pending"],
                   ["Compliance Officer",    PH,                    "_______________ / __________", "Pending"],
                   ["Head of Desk",          PH,                    "_______________ / __________", "Pending"],
               ] if sample else [[PH]*4 for _ in range(5)]),
               col_widths=[Cm(4.5), Cm(4), Cm(5.5), Cm(3)])

    # footer
    footer(doc, "Rey Capital — CONFIDENTIAL — CFD Strategy Spec v1.0  |  Do not distribute.")

    # metadata
    cp = doc.core_properties
    cp.title = "CFD Strategy Specification — Rey Capital"
    cp.author = "Rey Capital"
    cp.category = "CONFIDENTIAL"
    cp.keywords = "CONFIDENTIAL, CFD, Strategy, Rey Capital"
    cp.comments = "Confidential — Rey Capital internal use only."

    doc.save(path)
    print(f"Saved: {path}")


if __name__ == "__main__":
    build(f"{OUT}/CFD_Strategy_Template_v1.docx",                              sample=False)
    build(f"{OUT}/CFD_Strategy_Template_SAMPLE_EMA_9_21_Crossover.docx",       sample=True)
