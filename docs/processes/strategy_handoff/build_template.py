"""Build the Strategy Handoff & Go-Live Acceptance template (Rey Capital).

Two-stage gate document the manual trading team fills in before a strategy
moves into automation, and again before it goes live with real capital.

Run:
    python3 build_template.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Inches, Pt, RGBColor

# ---- Brand ---------------------------------------------------------------
BRAND_BLUE_HEX = "004AAC"
BRAND_LIGHT_HEX = "BDD4F5"
SOFT_HEX = "F4F7FC"
BRAND_BLUE = RGBColor(0x00, 0x4A, 0xAC)
DARK_TEXT = RGBColor(0x1A, 0x1A, 0x1A)
MUTED = RGBColor(0x6B, 0x72, 0x80)

FONT = "Calibri"

HERE = Path(__file__).resolve().parent
REPO_ROOT = HERE.parents[2]
LOGO = REPO_ROOT / "monitoring" / "dashboards" / "mt5_pnl_dashboard" / "src" / "TradeAtlas" / "Assets" / "ReyCapital_Logo.png"
OUTPUT = HERE / "Strategy_Handoff_Acceptance_Template.docx"


# ---- Helpers -------------------------------------------------------------

def shade_cell(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def set_cell_borders(cell, color_hex: str = "BFCBDB", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        b = OxmlElement(f"w:{edge}")
        b.set(qn("w:val"), "single")
        b.set(qn("w:sz"), size)
        b.set(qn("w:color"), color_hex)
        borders.append(b)
    tc_pr.append(borders)


def style_run(run, *, size: int = 11, bold: bool = False, color: RGBColor = DARK_TEXT) -> None:
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = color


def add_para(doc, text: str = "", *, size: int = 11, bold: bool = False, color=DARK_TEXT, align=None, space_after: int = 6):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        run = p.add_run(text)
        style_run(run, size=size, bold=bold, color=color)
    return p


def add_section_heading(doc, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    style_run(run, size=14, bold=True, color=BRAND_BLUE)
    # Bottom border on the heading paragraph for a divider effect.
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "8")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), BRAND_LIGHT_HEX)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_kv_table(doc, rows: list[tuple[str, str]], *, label_width_cm: float = 5.0, value_width_cm: float = 11.5) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for row_idx, (label, placeholder) in enumerate(rows):
        row = table.rows[row_idx]
        c1, c2 = row.cells
        c1.width = Cm(label_width_cm)
        c2.width = Cm(value_width_cm)
        shade_cell(c1, SOFT_HEX)
        set_cell_borders(c1)
        set_cell_borders(c2)
        for cell in (c1, c2):
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = c1.paragraphs[0]
        r1 = p1.add_run(label)
        style_run(r1, size=10, bold=True, color=BRAND_BLUE)
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(placeholder)
        style_run(r2, size=10, color=MUTED)


def add_checklist(doc, items: list[str]) -> None:
    table = doc.add_table(rows=len(items), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for idx, item in enumerate(items):
        row = table.rows[idx]
        c1, c2 = row.cells
        c1.width = Cm(1.2)
        c2.width = Cm(15.3)
        set_cell_borders(c1)
        set_cell_borders(c2)
        # Checkbox column.
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run("☐")
        style_run(r1, size=14, color=BRAND_BLUE)
        # Item text.
        p2 = c2.paragraphs[0]
        r2 = p2.add_run(item)
        style_run(r2, size=10)


def add_table_with_header(doc, headers: list[str], placeholder_rows: int, col_widths_cm: list[float]) -> None:
    table = doc.add_table(rows=placeholder_rows + 1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    for col_idx, header in enumerate(headers):
        cell = table.rows[0].cells[col_idx]
        cell.width = Cm(col_widths_cm[col_idx])
        shade_cell(cell, BRAND_BLUE_HEX)
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        r = p.add_run(header)
        style_run(r, size=10, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))
    for row_idx in range(1, placeholder_rows + 1):
        for col_idx in range(len(headers)):
            cell = table.rows[row_idx].cells[col_idx]
            cell.width = Cm(col_widths_cm[col_idx])
            set_cell_borders(cell)
            cell.text = ""


def add_header_band(doc) -> None:
    section = doc.sections[0]
    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if LOGO.exists():
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(1.6))
    p2 = header.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r2 = p2.add_run("Strategy Handoff & Go-Live Acceptance")
    style_run(r2, size=10, bold=True, color=BRAND_BLUE)


def add_footer(doc) -> None:
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Rey Capital  •  Smart Investments  •  Internal Use Only")
    style_run(r, size=9, color=MUTED)


# ---- Document content ----------------------------------------------------

def build() -> Path:
    doc = Document()

    # Page margins.
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

    add_header_band(doc)
    add_footer(doc)

    # Title block.
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    tr = title.add_run("Strategy Handoff & Go-Live Acceptance")
    style_run(tr, size=22, bold=True, color=BRAND_BLUE)

    sub = doc.add_paragraph()
    sr = sub.add_run("Manual Trading Team  →  Automation Team")
    style_run(sr, size=12, color=MUTED)

    add_para(
        doc,
        "This template is the single source of truth for handing off a profitable strategy from "
        "the manual trading team to the automation team. It defines two acceptance gates: "
        "Stage 1 — Handoff to Automation (the strategy is well-specified enough to be coded), "
        "and Stage 2 — Go-Live in Production (the automated build is proven and ready for real "
        "capital). A strategy may not advance to the next stage until every item in that stage's "
        "checklist is signed off by both teams.",
        size=10,
    )

    # ---- Strategy Identity ------------------------------------------------
    add_section_heading(doc, "1. Strategy Identity")
    add_kv_table(doc, [
        ("Strategy name", "<short, unique name>"),
        ("Author / submitter", "<manual team member>"),
        ("Submission date", "<YYYY-MM-DD>"),
        ("Version", "<v1.0>"),
        ("Strategy ID", "<assigned by automation team>"),
        ("Short description", "<one paragraph: what edge does this exploit?>"),
        ("Markets / instruments", "<e.g. NQ, ES, EURUSD, NIFTY>"),
        ("Timeframes", "<e.g. 5m chart, 1h confirmation>"),
        ("Trading sessions", "<e.g. NY 09:30–11:30 ET>"),
    ])

    # ---- Stage 1 ----------------------------------------------------------
    add_section_heading(doc, "Stage 1 — Handoff Acceptance Criteria")
    add_para(
        doc,
        "Goal: the strategy is unambiguous enough for an engineer to implement without further interpretation.",
        size=10,
        color=MUTED,
    )

    # 2. Indicators & Settings
    add_section_heading(doc, "2. Indicators & Settings")
    add_para(
        doc,
        "List every indicator the strategy depends on and its exact configuration. "
        "Anything not listed here will not be implemented.",
        size=10,
        color=MUTED,
    )
    add_table_with_header(
        doc,
        headers=["#", "Indicator", "Source / library", "Parameters", "Applied to (TF / price)", "Purpose"],
        placeholder_rows=6,
        col_widths_cm=[0.8, 3.0, 2.8, 4.0, 3.0, 2.9],
    )

    # 3. Strategy Logic
    add_section_heading(doc, "3. Strategy Logic Clarity")
    add_checklist(doc, [
        "Entry rules are written as deterministic, ordered conditions (no discretion words like \"usually\", \"sometimes\").",
        "Exit rules are fully specified: stop loss, take profit, time-based exit, trailing logic.",
        "Trade direction logic (long-only / short-only / both) is explicit.",
        "Confluence / filter conditions are listed in the order they are evaluated.",
        "Every parameter has a default value AND an allowed range.",
        "Edge cases documented: gaps, news, holidays, partial fills, missing data.",
        "A worked example trade is attached (chart + numbered steps).",
    ])
    add_para(doc, "Entry rules:", size=11, bold=True, space_after=2)
    add_para(doc, "<step-by-step entry conditions>", size=10, color=MUTED)
    add_para(doc, "Exit rules:", size=11, bold=True, space_after=2)
    add_para(doc, "<stop loss / take profit / trailing / time-based exits>", size=10, color=MUTED)
    add_para(doc, "Filters & confluences:", size=11, bold=True, space_after=2)
    add_para(doc, "<HTF bias, session, volatility, news avoidance, etc.>", size=10, color=MUTED)

    # 4. Risk & Money Management
    add_section_heading(doc, "4. Risk & Money Management")
    add_kv_table(doc, [
        ("Risk per trade", "<e.g. 0.5% of equity>"),
        ("Position sizing rule", "<fixed lot / % equity / volatility-based>"),
        ("Max concurrent positions", "<e.g. 2>"),
        ("Max daily loss", "<e.g. 2% / hard cutoff>"),
        ("Max drawdown stop", "<e.g. 8% pause-and-review>"),
        ("Correlation rules", "<no two longs on correlated instruments>"),
        ("Leverage cap", "<e.g. 1:10>"),
    ])

    # 5. Backtest Performance
    add_section_heading(doc, "5. Backtest Performance (Minimum Bar)")
    add_para(
        doc,
        "The strategy must clear ALL minimum thresholds below on a clean, out-of-sample backtest before handoff.",
        size=10,
        color=MUTED,
    )
    add_table_with_header(
        doc,
        headers=["Metric", "Minimum required", "Reported value", "Pass / Fail"],
        placeholder_rows=10,
        col_widths_cm=[5.5, 4.0, 4.0, 3.0],
    )
    add_para(doc, "Suggested metrics to populate above:", size=10, bold=True, space_after=2)
    add_para(
        doc,
        "Sample size (≥ 100 trades)  •  Win rate  •  Profit factor (≥ 1.5)  •  Expectancy per trade  •  "
        "Max drawdown (≤ 15%)  •  Sharpe ratio (≥ 1.0)  •  Avg R:R  •  Longest losing streak  •  "
        "Out-of-sample period covered  •  Walk-forward stability",
        size=10,
        color=MUTED,
    )

    # 6. Operational Readiness
    add_section_heading(doc, "6. Operational Readiness")
    add_kv_table(doc, [
        ("Broker(s) supported", "<e.g. ICMarkets, Zerodha, IBKR>"),
        ("Data feed assumed", "<broker feed / external>"),
        ("Spread assumption", "<e.g. ≤ 1.2 pips>"),
        ("Slippage assumption", "<e.g. 0.5 pip / 1 tick>"),
        ("Commission model", "<per lot / per side>"),
        ("Latency tolerance", "<e.g. ≤ 200ms>"),
        ("Required platform features", "<MT5, hedging on/off, OCO, etc.>"),
    ])

    # Stage 1 Sign-off
    add_section_heading(doc, "7. Stage 1 Sign-off")
    add_table_with_header(
        doc,
        headers=["Role", "Name", "Date", "Signature"],
        placeholder_rows=3,
        col_widths_cm=[5.0, 4.0, 3.5, 4.0],
    )
    add_para(
        doc,
        "Required signers: Manual team lead (submitter) • Automation team lead (acceptor) • Risk reviewer",
        size=9,
        color=MUTED,
    )

    # ---- Stage 2 ----------------------------------------------------------
    doc.add_page_break()
    add_section_heading(doc, "Stage 2 — Go-Live Acceptance Criteria")
    add_para(
        doc,
        "Goal: the automated implementation matches the spec, has been independently validated, "
        "and is operationally safe to run with real capital.",
        size=10,
        color=MUTED,
    )

    # 8. Implementation Verification
    add_section_heading(doc, "8. Implementation Verification")
    add_checklist(doc, [
        "Code is committed to the repo with a link to this document.",
        "Implementation reviewed by a second engineer (4-eyes).",
        "Unit tests cover all entry, exit, and filter conditions.",
        "Indicator outputs match the manual team's reference values on a held-out sample (≥ 50 bars).",
        "Worked example trade from Section 3 reproduces identically end-to-end in code.",
    ])

    # 9. Re-Backtest Parity
    add_section_heading(doc, "9. Re-Backtest Parity vs. Manual Spec")
    add_para(
        doc,
        "The automated backtest must reproduce the manual team's reported numbers within tolerance.",
        size=10,
        color=MUTED,
    )
    add_table_with_header(
        doc,
        headers=["Metric", "Manual reported", "Automated reproduced", "Delta", "Within tolerance?"],
        placeholder_rows=6,
        col_widths_cm=[4.0, 3.5, 3.5, 2.5, 3.0],
    )

    # 10. Forward Test / Demo
    add_section_heading(doc, "10. Forward Test on Demo")
    add_kv_table(doc, [
        ("Demo period start", "<YYYY-MM-DD>"),
        ("Demo period end", "<YYYY-MM-DD>"),
        ("Trades executed", "<count>"),
        ("Realized expectancy vs. backtest", "<within X% tolerance>"),
        ("Slippage observed vs. assumed", "<delta>"),
        ("Errors / interventions during run", "<list any>"),
    ])

    # 11. Operational Safeguards
    add_section_heading(doc, "11. Operational Safeguards")
    add_checklist(doc, [
        "Kill-switch / pause command tested and documented.",
        "Daily-loss cutoff implemented and alert-tested.",
        "Position-size cap enforced in code (not just config).",
        "Logging covers every order, fill, rejection, and error with timestamps.",
        "Monitoring dashboard shows the strategy P&L, open positions, and last heartbeat.",
        "On-call runbook written: who responds, what to check, how to halt.",
        "Disaster scenarios rehearsed: broker disconnect, data outage, server restart.",
        "Capital allocation for go-live agreed and documented.",
    ])

    # 12. Go-Live Decision
    add_section_heading(doc, "12. Go-Live Decision")
    add_kv_table(doc, [
        ("Initial capital allocation", "<amount or %>"),
        ("Ramp-up plan", "<e.g. 25% → 50% → 100% over 4 weeks>"),
        ("Review checkpoint date", "<YYYY-MM-DD>"),
        ("Decommission criteria", "<e.g. drawdown > 10% or 20 losing trades>"),
    ])

    # 13. Stage 2 Sign-off
    add_section_heading(doc, "13. Stage 2 Sign-off")
    add_table_with_header(
        doc,
        headers=["Role", "Name", "Date", "Signature"],
        placeholder_rows=4,
        col_widths_cm=[5.0, 4.0, 3.5, 4.0],
    )
    add_para(
        doc,
        "Required signers: Manual team lead • Automation team lead • Risk reviewer • Head of Trading",
        size=9,
        color=MUTED,
    )

    # ---- Appendix ---------------------------------------------------------
    doc.add_page_break()
    add_section_heading(doc, "Appendix A — Attachments Checklist")
    add_checklist(doc, [
        "Strategy logic document (PDF or DOCX).",
        "Backtest report (full equity curve, trade list, drawdown chart).",
        "Out-of-sample / walk-forward report.",
        "Annotated chart screenshots showing 3+ example trades.",
        "Indicator source code or formulas.",
        "Risk model spreadsheet.",
    ])

    add_section_heading(doc, "Appendix B — Revision History")
    add_table_with_header(
        doc,
        headers=["Version", "Date", "Author", "Change summary"],
        placeholder_rows=5,
        col_widths_cm=[2.0, 3.0, 4.0, 7.5],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    path = build()
    print(f"Wrote {path}")
