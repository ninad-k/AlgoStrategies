"""
Generate professional Word documents for Rey Capital AI Bot and ReySentinel.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE


def set_cell_shading(cell, color_hex):
    """Set cell background color."""
    from docx.oxml.ns import qn
    from lxml import etree
    shading = etree.SubElement(cell._tc.get_or_add_tcPr(), qn('w:shd'))
    shading.set(qn('w:fill'), color_hex)
    shading.set(qn('w:val'), 'clear')


def add_styled_table(doc, headers, rows, col_widths=None):
    """Add a professionally styled table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.name = 'Arial'
        set_cell_shading(cell, '0052CC')

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, value in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(value)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Arial'
            if row_idx % 2 == 1:
                set_cell_shading(cell, 'F0F4F8')

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    return table


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x00, 0x52, 0xCC)
    return h


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Arial'
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Arial'
    else:
        run = p.add_run(text)
        run.font.size = Pt(10)
        run.font.name = 'Arial'
    return p


def add_body(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.name = 'Arial'
    p.paragraph_format.space_after = Pt(6)
    return p


# ════════════════════════════════════════════════════════════
# DOCUMENT 1: Rey Capital AI Bot
# ════════════════════════════════════════════════════════════

def create_rey_capital_ai_bot_doc():
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # ── Cover Page ──
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('Rey Capital AI Bot')
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x52, 0xCC)
    run.font.name = 'Arial'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Architecture, Operations & Roadmap')
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor(0x5A, 0x63, 0x70)
    run.font.name = 'Arial'

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Rey Capital | Confidential')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)

    doc.add_page_break()

    # ── Table of Contents ──
    add_heading_styled(doc, 'Table of Contents', level=1)
    toc_items = [
        '1. Executive Summary',
        '2. System Architecture',
        '3. Technical Components',
        '4. Signal Flow & Decision Pipeline',
        '5. Risk Management Framework',
        '6. Self-Learning & Adaptation',
        '7. Dashboard & Monitoring',
        '8. Deployment & Infrastructure',
        '9. Current Capabilities',
        '10. Roadmap & Next Steps',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.name = 'Arial'
        run.font.color.rgb = RGBColor(0x00, 0x52, 0xCC)

    doc.add_page_break()

    # ── 1. Executive Summary ──
    add_heading_styled(doc, '1. Executive Summary')
    add_body(doc, 'Rey Capital AI Bot is a fully automated cryptocurrency scalping system powered by Google Gemma 4 AI (via Ollama). It operates on 1-minute candles across 5 crypto CFD pairs on MetaTrader 5, analyzing 30+ technical indicators per cycle and making autonomous trade decisions with built-in risk management.')

    add_body(doc, 'The system features a self-learning adaptive engine that reviews its own trade outcomes, adjusts confidence thresholds, and feeds lessons back into future decisions. A professional web dashboard provides real-time monitoring with dark/light theme support.')

    add_heading_styled(doc, 'Key Metrics', level=2)
    add_styled_table(doc,
        ['Metric', 'Value'],
        [
            ['Trading Pairs', 'BTCUSD, ETHUSD, LTCUSD, XRPUSD, SOLUSD'],
            ['Timeframe', '1-minute candles (scalping)'],
            ['Indicators Analyzed', '30+ per cycle per symbol'],
            ['AI Model', 'Gemma 4 via Ollama (local inference)'],
            ['Risk Per Trade', '1% of account balance'],
            ['Max Open Trades', '5 concurrent positions'],
            ['Daily Loss Limit', '5% of account'],
            ['Execution', 'MetaTrader 5 (live) or Paper mode'],
        ],
        col_widths=[2.5, 4.0]
    )

    doc.add_page_break()

    # ── 2. System Architecture ──
    add_heading_styled(doc, '2. System Architecture')
    add_body(doc, 'The bot follows a modular pipeline architecture with clear separation of concerns:')

    add_styled_table(doc,
        ['Layer', 'Component', 'Responsibility'],
        [
            ['Data', 'MT5 Data Feed', 'Fetches 500 bars of OHLCV per symbol from MetaTrader 5'],
            ['Data', 'TradingView Fallback', 'Scanner API fallback when MT5 is unavailable'],
            ['Analysis', 'Indicator Engine', '30+ technical indicators via pandas_ta'],
            ['AI', 'Gemma 4 Analyzer', 'LLM-based trade decision with JSON response'],
            ['Risk', 'Risk Manager', '7-point validation, position sizing, cooldowns'],
            ['Execution', 'Broker Bridge', 'MT5/Binance/Paper order execution'],
            ['Learning', 'Trade Reviewer', 'Self-analysis, adaptive context generation'],
            ['UI', 'Dashboard', 'Flask + SocketIO real-time web interface'],
        ],
        col_widths=[1.0, 2.0, 3.5]
    )

    add_heading_styled(doc, 'Architecture Diagram', level=2)
    add_body(doc, 'MT5 Terminal (candles)')
    add_body(doc, '    \u2193')
    add_body(doc, 'MT5DataFeed.get_candles() \u2192 500 bars OHLCV')
    add_body(doc, '    \u2193')
    add_body(doc, 'calculate_indicators() \u2192 30+ indicators')
    add_body(doc, '    \u2193')
    add_body(doc, 'analyze_with_gemma() \u2192 Ollama/Gemma 4 HTTP')
    add_body(doc, '    \u2193')
    add_body(doc, 'RiskManager.can_trade() \u2192 7-point validation')
    add_body(doc, '    \u2193')
    add_body(doc, 'RiskManager.calculate_position_size()')
    add_body(doc, '    \u2193')
    add_body(doc, 'BrokerBridge.place_order() \u2192 MT5 execution')
    add_body(doc, '    \u2193')
    add_body(doc, 'Trade Journal + Outcome Tracking + Self-Learning')

    doc.add_page_break()

    # ── 3. Technical Components ──
    add_heading_styled(doc, '3. Technical Components')

    add_heading_styled(doc, '3.1 Indicator Engine (30+ Indicators)', level=2)
    add_styled_table(doc,
        ['Category', 'Indicators'],
        [
            ['Trend (9)', 'EMA(9,20,50,200), SMA(20,50), ADX, Supertrend, Parabolic SAR'],
            ['Ichimoku (5)', 'Tenkan-sen, Kijun-sen, Span A/B, Cloud Color, Signal'],
            ['Momentum (7)', 'RSI(14), MACD(12,26,9), Stoch RSI, CCI(20), Williams %R, ROC(10), MFI(14)'],
            ['Volatility (4)', 'ATR(14), Bollinger Bands(20,2), Keltner Channels, Donchian'],
            ['Volume (5)', 'Current vs 20-SMA ratio, VWAP, OBV, A/D, Volume Trend'],
            ['Price Action (3)', 'Last 5 candles, Support/Resistance (swing points), Candlestick Patterns'],
            ['Patterns (10+)', 'Doji, Hammer, Shooting Star, Engulfing, Morning/Evening Star, etc.'],
        ],
        col_widths=[1.5, 5.0]
    )

    add_heading_styled(doc, '3.2 Gemma 4 AI Analyzer', level=2)
    add_body(doc, 'The AI analyzer sends all 30+ indicators as a structured prompt to Gemma 4 via Ollama HTTP API. The system prompt encodes crypto-specific trading rules, instrument characteristics, and risk discipline. Key configuration:')
    add_bullet(doc, 'Temperature: 0.1 (low for reproducible decisions)')
    add_bullet(doc, 'Max tokens: 8,192')
    add_bullet(doc, 'Response format: JSON with action, confidence, SL/TP ATR multipliers, reason')
    add_bullet(doc, 'Adaptive context from past trades is prepended to every prompt')

    add_heading_styled(doc, '3.3 Broker Bridge', level=2)
    add_body(doc, 'Supports three execution modes:')
    add_bullet(doc, 'MT5 Broker: ', bold_prefix='Live Trading: ')
    add_bullet(doc, 'Binance via CCXT for crypto futures', bold_prefix='Crypto Exchange: ')
    add_bullet(doc, 'In-memory simulation for testing', bold_prefix='Paper Mode: ')
    add_body(doc, 'Auto-detects filling mode (FOK/IOC/RETURN) per symbol with retry logic.')

    doc.add_page_break()

    # ── 4. Signal Flow ──
    add_heading_styled(doc, '4. Signal Flow & Decision Pipeline')
    add_body(doc, 'Each 60-second cycle processes all 5 symbols sequentially:')

    steps = [
        ('Fetch Data', 'Pull 500 bars from MT5 (or TradingView fallback)'),
        ('Calculate Indicators', 'Compute 30+ technical indicators via pandas_ta'),
        ('AI Analysis', 'Send indicators to Gemma 4, receive JSON decision'),
        ('Risk Validation', '7-point check: symbol, confidence, max trades, duplicates, daily loss, cooldown, streak'),
        ('Position Sizing', 'Calculate lots from ATR, tick value, and 1% risk rule'),
        ('Price Adjustment', 'Get real-time bid/ask for spread-aware SL/TP placement'),
        ('Order Execution', 'Place order on MT5 with magic number 240411'),
        ('Registration', 'Log trade, update cooldown timer, write journal entry'),
        ('WebSocket Update', 'Push real-time update to dashboard via SocketIO'),
    ]
    add_styled_table(doc,
        ['Step', 'Action', 'Details'],
        [(str(i+1), s, d) for i, (s, d) in enumerate(steps)],
        col_widths=[0.5, 1.8, 4.2]
    )

    doc.add_page_break()

    # ── 5. Risk Management ──
    add_heading_styled(doc, '5. Risk Management Framework')

    add_heading_styled(doc, '5.1 Seven-Point Risk Check', level=2)
    checks = [
        ('Symbol Whitelist', 'Symbol must be in allowed_symbols list'),
        ('Confidence Threshold', 'AI confidence >= dynamic threshold (0.50-0.85)'),
        ('Max Open Trades', 'Must not exceed max_open_trades (default: 5)'),
        ('No Duplicate Position', 'Cannot have two positions on same symbol'),
        ('Daily Loss Limit', 'Daily P&L must be above -5% of account'),
        ('Per-Symbol Cooldown', 'Minimum 3 minutes between trades on same symbol'),
        ('Streak Cooldown', 'After 3 consecutive losses on a symbol: 15-min cooldown'),
    ]
    add_styled_table(doc,
        ['#', 'Check', 'Rule'],
        [(str(i+1), c, r) for i, (c, r) in enumerate(checks)],
        col_widths=[0.4, 2.0, 4.1]
    )

    add_heading_styled(doc, '5.2 Position Sizing Formula', level=2)
    add_body(doc, 'Uses MT5 symbol specifications for precise lot calculation:')
    add_body(doc, 'risk_amount = account_balance x (max_position_size_pct / 100)')
    add_body(doc, 'sl_distance = ATR x sl_atr_multiplier')
    add_body(doc, 'ticks_in_sl = sl_distance / tick_size')
    add_body(doc, 'lots = risk_amount / (ticks_in_sl x tick_value)')
    add_body(doc, 'lots = clamp(lots, min_lot, max_lot) rounded to lot_step')

    doc.add_page_break()

    # ── 6. Self-Learning ──
    add_heading_styled(doc, '6. Self-Learning & Adaptation')
    add_body(doc, 'The bot continuously improves through a multi-layered feedback loop:')

    add_heading_styled(doc, '6.1 Performance Analysis', level=2)
    add_bullet(doc, 'Triggered every N trades (configurable, default: 3)')
    add_bullet(doc, 'Calculates: win rate, avg win/loss, profit factor, per-symbol breakdown')
    add_bullet(doc, 'Evaluates indicator effectiveness: RSI zones, trend alignment, Ichimoku signals, volume levels')

    add_heading_styled(doc, '6.2 Adaptive Context Generation', level=2)
    add_body(doc, 'Builds human-readable lessons that are prepended to every Gemma prompt:')
    add_bullet(doc, 'Overall stats: win rate, avg win/loss, profit factor')
    add_bullet(doc, 'Per-symbol biases: "FAVOR BTCUSD BUY signals (78% win rate)"')
    add_bullet(doc, 'Indicator patterns: "RSI oversold: 68% win rate - FAVORABLE"')

    add_heading_styled(doc, '6.3 Dynamic Threshold Adjustment', level=2)
    add_bullet(doc, 'Win rate < 40%: raise confidence threshold by +0.05 (be more selective)')
    add_bullet(doc, 'Win rate > 60%: lower threshold by -0.02 (take more trades)')
    add_bullet(doc, 'Range clamped between 0.50 and 0.85')

    add_heading_styled(doc, '6.4 Weekly Meta-Review', level=2)
    add_body(doc, 'Gemma reviews its last 50 trades and generates meta-patterns: which indicator combinations work, which to avoid, and market condition insights.')

    doc.add_page_break()

    # ── 7. Dashboard ──
    add_heading_styled(doc, '7. Dashboard & Monitoring')
    add_body(doc, 'Professional web dashboard served on port 8050 with Flask + SocketIO:')

    add_styled_table(doc,
        ['Feature', 'Description'],
        [
            ['Real-time Updates', 'WebSocket push for decisions, trades, and closures'],
            ['Symbol Cards', 'Live price, indicators, AI decision per symbol'],
            ['Decision Timeline', 'Chronological log of all Gemma decisions'],
            ['Trade Journal', 'Full AI reasoning, strategy classification, risk details'],
            ['AI Learning Panel', 'Adaptive context, win rate, parameter adjustments'],
            ['Lot Size Controls', 'Manual override per symbol from dashboard'],
            ['Dark/Light Theme', 'Toggle with Rey Capital branding (black/white)'],
            ['User Comments', 'Add retrospective notes to trade journal entries'],
        ],
        col_widths=[2.0, 4.5]
    )

    doc.add_page_break()

    # ── 8. Deployment ──
    add_heading_styled(doc, '8. Deployment & Infrastructure')

    add_heading_styled(doc, '8.1 Requirements', level=2)
    add_bullet(doc, 'Windows PC with MetaTrader 5 desktop running')
    add_bullet(doc, 'Ollama with Gemma 4 model installed locally')
    add_bullet(doc, 'Python 3.11+ with dependencies (flask, pandas_ta, requests, etc.)')

    add_heading_styled(doc, '8.2 Deployment Options', level=2)
    add_styled_table(doc,
        ['Environment', 'Setup', 'Notes'],
        [
            ['Local Windows', 'setup_local.bat', 'MT5 + Ollama on same machine'],
            ['AWS EC2 Windows', 'setup_windows_ec2.ps1', 'RDP access, MT5 + Ollama on EC2'],
            ['Paper Mode', 'python run.py --mode paper', 'No broker needed, simulated execution'],
        ],
        col_widths=[1.5, 2.0, 3.0]
    )

    add_heading_styled(doc, '8.3 Startup Command', level=2)
    add_body(doc, 'python run.py --config config.yaml --port 8050 --mode paper')

    doc.add_page_break()

    # ── 9. Current Capabilities ──
    add_heading_styled(doc, '9. Current Capabilities')
    add_styled_table(doc,
        ['Capability', 'Status', 'Details'],
        [
            ['Crypto Scalping', 'LIVE', '5 pairs on 1M with Gemma 4'],
            ['30+ Indicators', 'LIVE', 'Ichimoku, Supertrend, MACD, RSI, BB, ADX, CCI, etc.'],
            ['Self-Learning', 'LIVE', 'Adaptive context, threshold adjustment, weekly review'],
            ['Trade Journal', 'LIVE', 'Full AI reasoning, strategy classification, comments'],
            ['Risk Management', 'LIVE', '7-point check, position sizing, cooldowns, streak protection'],
            ['Web Dashboard', 'LIVE', 'Real-time Flask+SocketIO with dark/light theme'],
            ['Paper Trading', 'LIVE', 'Full simulation without real money'],
            ['MT5 Execution', 'LIVE', 'Auto-detect filling mode, retry logic'],
            ['Binance', 'AVAILABLE', 'CCXT integration for crypto futures'],
            ['AWS Deployment', 'AVAILABLE', 'EC2 setup scripts included'],
        ],
        col_widths=[1.8, 1.0, 3.7]
    )

    doc.add_page_break()

    # ── 10. Roadmap ──
    add_heading_styled(doc, '10. Roadmap & Next Steps')

    add_heading_styled(doc, 'Short-Term (1-2 weeks)', level=2)
    add_bullet(doc, 'Integrate with ReySentinel multi-model ensemble for higher accuracy')
    add_bullet(doc, 'Add market regime detection to filter trades by market condition')
    add_bullet(doc, 'Implement sentiment overlay from ReySentinel for directional bias')

    add_heading_styled(doc, 'Medium-Term (1-2 months)', level=2)
    add_bullet(doc, 'Expand to forex and commodity pairs (XAUUSD, EURUSD, etc.)')
    add_bullet(doc, 'Add volume profile integration for POC/VAH/VAL level awareness')
    add_bullet(doc, 'Build mobile push notifications for trade alerts')
    add_bullet(doc, 'Implement trailing stop optimization based on ATR')

    add_heading_styled(doc, 'Long-Term (3-6 months)', level=2)
    add_bullet(doc, 'Multi-account support: run bot across 10+ MT5 accounts')
    add_bullet(doc, 'Cloud deployment with Docker/Kubernetes for horizontal scaling')
    add_bullet(doc, 'Advanced pattern recognition (CNN) for chart pattern detection')
    add_bullet(doc, 'Compliance-grade audit logging for regulatory requirements')
    add_bullet(doc, 'Mobile trading app with one-click override capabilities')

    # Save
    path = os.path.join(os.path.dirname(__file__), 'Rey_Capital_AI_Bot_Architecture.docx')
    doc.save(path)
    print(f'Created: {path}')
    return path


# ════════════════════════════════════════════════════════════
# DOCUMENT 2: ReySentinel
# ════════════════════════════════════════════════════════════

def create_reysentinel_doc():
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(10)

    # ── Cover Page ──
    for _ in range(6):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run('ReySentinel')
    run.font.size = Pt(36)
    run.bold = True
    run.font.color.rgb = RGBColor(0x00, 0x52, 0xCC)
    run.font.name = 'Arial'

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run('Trading Intelligence Platform')
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x5A, 0x63, 0x70)
    run.font.name = 'Arial'

    doc.add_paragraph()
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub2.add_run('Architecture, Operations & Roadmap')
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)

    doc.add_paragraph()
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run('Rey Capital | Confidential')
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x8B, 0x94, 0x9E)

    doc.add_page_break()

    # ── TOC ──
    add_heading_styled(doc, 'Table of Contents', level=1)
    toc = [
        '1. Executive Summary',
        '2. Platform Overview',
        '3. System Architecture',
        '4. Module Deep Dives',
        '5. Signal Flow: End-to-End',
        '6. Dashboards & UI',
        '7. Deployment Options',
        '8. Mobile App',
        '9. How It All Works Together',
        '10. Current Status',
        '11. Roadmap & What We Can Build Next',
    ]
    for item in toc:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x00, 0x52, 0xCC)

    doc.add_page_break()

    # ── 1. Executive Summary ──
    add_heading_styled(doc, '1. Executive Summary')
    add_body(doc, 'ReySentinel is a self-contained trading intelligence platform that combines multiple AI models, real-time market analysis, and professional dashboards into a unified system. It extends the Rey Capital AI Bot with multi-model consensus, market regime detection, sentiment analysis, volume profiling, correlation tracking, and pattern recognition.')

    add_body(doc, 'The platform is designed to be copied as a single folder with all dependencies included. It runs on any machine with Python 3.11+ and can be deployed to cloud infrastructure via Docker/Kubernetes.')

    add_heading_styled(doc, 'Platform at a Glance', level=2)
    add_styled_table(doc,
        ['Metric', 'Value'],
        [
            ['Total Modules', '11 (AI, analysis, dashboards, deployment, mobile)'],
            ['Total Files', '84+ production-ready files'],
            ['AI Models', 'Gemma 4 + LLaMA 3 + ONNX/LightGBM ensemble'],
            ['Dashboards', '2 web dashboards (Heatmap :8061, Multi-Account :8062)'],
            ['Mobile', 'React Native (Expo) for iOS/Android'],
            ['Deployment', 'Docker + Kubernetes + one-click deploy scripts'],
            ['Self-Contained', 'Copy folder + pip install + run'],
        ],
        col_widths=[2.0, 4.5]
    )

    doc.add_page_break()

    # ── 2. Platform Overview ──
    add_heading_styled(doc, '2. Platform Overview')
    add_body(doc, 'ReySentinel consists of 11 integrated modules organized into 4 tiers:')

    add_styled_table(doc,
        ['Tier', 'Module', 'Purpose'],
        [
            ['AI & Intelligence', 'Multi-Model Ensemble', 'Gemma 4 + LLaMA 3 + ONNX with weighted voting'],
            ['AI & Intelligence', 'Market Regime Detector', 'Classify trending/ranging/volatile/breakout states'],
            ['AI & Intelligence', 'Sentiment Bridge', 'Reddit + RSS + Fear&Greed \u2192 directional signal'],
            ['Data & Research', 'Volume Profile', 'Real-time TPO, POC, VAH/VAL, HVN/LVN'],
            ['Data & Research', 'Correlation Engine', 'Rolling correlation, lead-lag, cointegration pairs'],
            ['Data & Research', 'Pattern Recognition', 'CNN-based chart pattern detection (12 patterns)'],
            ['Risk & Compliance', 'Heatmap Dashboard', 'Portfolio exposure, VaR (3 methods), stress testing'],
            ['Risk & Compliance', 'Audit Logger', 'Compliance logging, backtest reconciliation, CSV export'],
            ['Infrastructure', 'Multi-Account', 'N MT5 accounts, consolidated P&L, risk allocation'],
            ['Infrastructure', 'Cloud Deployment', 'Docker, Kubernetes, deploy scripts'],
            ['Infrastructure', 'Mobile App', 'React Native with P&L charts, quick trade actions'],
        ],
        col_widths=[1.5, 2.0, 3.0]
    )

    doc.add_page_break()

    # ── 3. System Architecture ──
    add_heading_styled(doc, '3. System Architecture')

    add_heading_styled(doc, '3.1 Entry Point (app.py)', level=2)
    add_body(doc, 'The unified entry point launches all services as background threads:')
    add_bullet(doc, 'Trading Engine thread: data fetch \u2192 indicators \u2192 regime \u2192 ensemble \u2192 risk \u2192 execution')
    add_bullet(doc, 'Heatmap Dashboard thread: FastAPI on port 8061')
    add_bullet(doc, 'Multi-Account Dashboard thread: FastAPI on port 8062')
    add_bullet(doc, 'Audit Logger: background logging to JSONL + SQLite')

    add_heading_styled(doc, '3.2 Shared Layer', level=2)
    add_body(doc, 'All modules share common utilities copied from the Rey Capital AI Bot:')
    add_styled_table(doc,
        ['Component', 'File', 'Source'],
        [
            ['30+ Indicators', 'shared/indicators.py', 'Adapted from gemma_trader/local_trader.py'],
            ['MT5 Connector', 'shared/mt5_connector.py', 'Adapted from gemma_trader/mt5_data_feed.py'],
            ['Broker Abstraction', 'shared/broker.py', 'Adapted from gemma_trader/broker_bridge.py'],
            ['Risk Manager', 'shared/risk_manager.py', 'Adapted from gemma_trader/risk_manager.py'],
            ['Database', 'shared/db.py', 'SQLite helper for audit/decisions/outcomes'],
            ['Pydantic Models', 'shared/models.py', 'TradeDecision, EnsembleDecision, etc.'],
        ],
        col_widths=[1.5, 2.5, 2.5]
    )

    add_heading_styled(doc, '3.3 Data Flow', level=2)
    add_body(doc, 'MT5 Terminal \u2192 MT5 Connector (500 bars)')
    add_body(doc, '    \u2193')
    add_body(doc, 'Indicator Engine (30+ indicators)')
    add_body(doc, '    \u2193 (parallel)')
    add_body(doc, 'Regime Detector + Volume Profile + Pattern Recognition')
    add_body(doc, '    \u2193')
    add_body(doc, 'Ensemble Engine (Gemma + LLaMA + ONNX \u2192 weighted vote)')
    add_body(doc, '    \u2193')
    add_body(doc, 'Risk Manager (7-point check + position sizing)')
    add_body(doc, '    \u2193')
    add_body(doc, 'Broker Bridge (MT5/Binance/Paper execution)')
    add_body(doc, '    \u2193')
    add_body(doc, 'Audit Logger + Calibrator + Dashboard updates')

    doc.add_page_break()

    # ── 4. Module Deep Dives ──
    add_heading_styled(doc, '4. Module Deep Dives')

    # 4.1 Ensemble
    add_heading_styled(doc, '4.1 Multi-Model Ensemble Engine', level=2)
    add_body(doc, 'Three AI models analyze the same market data independently, then vote:')

    add_styled_table(doc,
        ['Model', 'Personality', 'Weight', 'Speed'],
        [
            ['Gemma 4', 'Aggressive scalper \u2014 takes trades with 2+ signals', '40%', '~5-10s'],
            ['LLaMA 3', 'Conservative analyst \u2014 needs 3+ signals, ADX>25', '35%', '~5-10s'],
            ['ONNX/LightGBM', 'Fast rule-based \u2014 indicator scoring, sub-second', '25%', '<100ms'],
        ],
        col_widths=[1.5, 3.0, 0.8, 1.2]
    )

    add_body(doc, 'Voting Methods:')
    add_bullet(doc, 'Weighted average of confidences, majority action wins', bold_prefix='Weighted Average: ')
    add_bullet(doc, 'Simple 2-of-3 majority', bold_prefix='Majority Vote: ')
    add_bullet(doc, 'Any model with HOLD at 85%+ confidence vetoes the trade', bold_prefix='Veto: ')

    add_body(doc, 'Dynamic Weight Adjustment: After 5+ trades, model weights shift based on individual win rates. A model that consistently wins gets more influence.')

    add_body(doc, 'Confidence Calibrator: Tracks predicted vs actual win rates per model per confidence bin. If Gemma says 0.8 confidence but only wins 50% at that level, calibrated output becomes ~0.5.')

    # 4.2 Regime
    add_heading_styled(doc, '4.2 Market Regime Detector', level=2)
    add_body(doc, 'Classifies market state every cycle using ATR ratio, ADX, BB width, volume, and return volatility:')

    add_styled_table(doc,
        ['Regime', 'Conditions', 'Trading Implication'],
        [
            ['TRENDING_UP', 'ADX>25, directional bias>+1.5, consistency>70%', 'Buy the dips, trail stops'],
            ['TRENDING_DOWN', 'ADX>25, directional bias<-1.5, consistency>70%', 'Sell the rallies'],
            ['RANGING', 'ADX<20, low volatility', 'Mean reversion, tight S/R'],
            ['VOLATILE', 'ATR ratio>1.5, high return std', 'Reduce size, wider stops'],
            ['BREAKOUT', 'ATR ratio>1.8 AND volume>2x', 'Ride momentum, aggressive'],
        ],
        col_widths=[1.5, 2.8, 2.2]
    )

    # 4.3 Sentiment
    add_heading_styled(doc, '4.3 Sentiment-to-Signal Bridge', level=2)
    add_body(doc, 'Scrapes 3 data sources, scores via LLM, and generates a -1.0 to +1.0 directional signal:')
    add_bullet(doc, 'r/cryptocurrency, r/Bitcoin, r/ethtrader (via PRAW)', bold_prefix='Reddit: ')
    add_bullet(doc, 'CoinTelegraph, CoinDesk RSS feeds', bold_prefix='RSS News: ')
    add_bullet(doc, 'Crypto Fear & Greed Index (alternative.me API)', bold_prefix='Fear & Greed: ')
    add_body(doc, 'Weighted aggregation: RSS 40% + Reddit 35% + Fear&Greed 25%')

    doc.add_page_break()

    # 4.4 Volume Profile
    add_heading_styled(doc, '4.4 Intraday Volume Profile', level=2)
    add_body(doc, 'Builds Time-Price-Opportunity (TPO) charts from live data:')
    add_bullet(doc, 'POC (Point of Control): highest-volume price level')
    add_bullet(doc, 'VAH/VAL: Value Area High/Low (70% of volume)')
    add_bullet(doc, 'HVN: High Volume Nodes (support/resistance)')
    add_bullet(doc, 'LVN: Low Volume Nodes (breakout targets)')
    add_body(doc, 'Updates in real-time as new candles arrive. Key levels are injected into ensemble decisions.')

    # 4.5 Correlation
    add_heading_styled(doc, '4.5 Correlation & Causality Engine', level=2)
    add_body(doc, 'Tracks relationships between trading pairs:')
    add_bullet(doc, 'Rolling Pearson and Spearman correlation (configurable window)')
    add_bullet(doc, 'Lead-lag detection via cross-correlation (which asset leads by how many bars)')
    add_bullet(doc, 'Engle-Granger cointegration test for mean-reversion pair suggestions')
    add_bullet(doc, 'Correlation break alerts when relationships change significantly')

    # 4.6 Pattern Recognition
    add_heading_styled(doc, '4.6 Pattern Recognition AI', level=2)
    add_body(doc, 'CNN-based chart pattern detection with 12 patterns:')

    add_styled_table(doc,
        ['Bullish', 'Bearish', 'Neutral'],
        [
            ['Double Bottom', 'Double Top', 'No Pattern'],
            ['Inv Head & Shoulders', 'Head & Shoulders', ''],
            ['Ascending Triangle', 'Descending Triangle', ''],
            ['Bull Flag', 'Bear Flag', ''],
            ['Cup & Handle', 'Rising Wedge', ''],
            ['Falling Wedge', '', ''],
        ],
        col_widths=[2.2, 2.2, 2.1]
    )

    add_body(doc, 'Architecture: Conv2d(1,32) \u2192 Conv2d(32,64) \u2192 Conv2d(64,128) \u2192 FC(512) \u2192 FC(12 classes)')
    add_body(doc, 'Falls back to rule-based swing point detection when no trained model is available.')

    doc.add_page_break()

    # 4.7 Audit Logger
    add_heading_styled(doc, '4.7 Trade Audit Logger', level=2)
    add_body(doc, 'Compliance-grade trade logging with dual storage:')
    add_bullet(doc, 'JSONL file: append-only, human-readable, every trade + decision + risk check')
    add_bullet(doc, 'SQLite database: structured queries, analytics, dashboarding')
    add_bullet(doc, 'Backtest reconciler: compare expected vs actual P&L')
    add_bullet(doc, 'CSV export: MiFID II / SEC compatible columns')

    # 4.8 Multi-Account
    add_heading_styled(doc, '4.8 Multi-Account Aggregator', level=2)
    add_body(doc, 'Manage 10+ MT5 accounts from a single dashboard:')
    add_bullet(doc, 'Connect to N MT5 terminals simultaneously')
    add_bullet(doc, 'Consolidated P&L across all accounts')
    add_bullet(doc, 'Risk allocation: proportional by equity or manual weights')
    add_bullet(doc, 'Per-account and per-symbol P&L breakdown')

    doc.add_page_break()

    # ── 5. Signal Flow ──
    add_heading_styled(doc, '5. Signal Flow: End-to-End')
    add_body(doc, 'Here is the complete decision pipeline for a single trading cycle:')

    flow = [
        ('1. Data Ingestion', 'MT5 Connector fetches 500 bars of 1M OHLCV for each of 5 symbols'),
        ('2. Indicator Computation', '30+ indicators calculated via pandas_ta (trend, momentum, volatility, volume, Ichimoku, etc.)'),
        ('3. Regime Detection', 'ATR ratio, ADX, BB width, volume analyzed. Market classified as TRENDING/RANGING/VOLATILE/BREAKOUT'),
        ('4. Volume Profile', 'TPO profile updated. POC, VAH, VAL levels injected into indicator data'),
        ('5. Sentiment Check', 'Cached sentiment signal (-1 to +1) from last scrape is attached (scraped every 15 min)'),
        ('6. Pattern Scan', 'CNN or rule-based detector identifies chart patterns in latest 64-bar window'),
        ('7. Ensemble Voting', 'Gemma 4, LLaMA 3, and ONNX model each produce BUY/SELL/HOLD with confidence. Weighted vote determines final action'),
        ('8. Calibration', 'Final confidence adjusted based on historical accuracy per model per confidence bin'),
        ('9. Risk Validation', '7-point check: symbol, confidence, max trades, duplicates, daily loss, cooldown, streak'),
        ('10. Position Sizing', 'Lots calculated from 1% risk rule using ATR and MT5 symbol specs'),
        ('11. Execution', 'Order placed on MT5/Binance/Paper with SL/TP from ATR multipliers'),
        ('12. Logging', 'Trade logged to audit trail, journal, SQLite. Model outcomes tracked for weight adjustment'),
    ]
    add_styled_table(doc,
        ['Step', 'Description'],
        [(s, d) for s, d in flow],
        col_widths=[1.8, 4.7]
    )

    doc.add_page_break()

    # ── 6. Dashboards ──
    add_heading_styled(doc, '6. Dashboards & UI')

    add_heading_styled(doc, '6.1 Portfolio Heatmap Dashboard (Port 8061)', level=2)
    add_styled_table(doc,
        ['Feature', 'Description'],
        [
            ['D3.js Treemap', 'Visual portfolio heatmap colored by P&L (green=profit, red=loss)'],
            ['VaR Gauges', 'Historical, Parametric, and Monte Carlo VaR at 95% confidence'],
            ['Stress Testing', '6 predefined scenarios: flash crash, rate shock, crypto winter, etc.'],
            ['Risk Warnings', 'Automatic alerts for concentration risk, high VaR, stress failures'],
            ['Dark/Light Theme', 'Toggle with Rey Capital logo swap'],
        ],
        col_widths=[2.0, 4.5]
    )

    add_heading_styled(doc, '6.2 Multi-Account Dashboard (Port 8062)', level=2)
    add_styled_table(doc,
        ['Feature', 'Description'],
        [
            ['Account Cards', 'Balance, equity, profit, leverage per account with status badges'],
            ['Risk Allocations', 'Per-account risk weight and dollar allocation'],
            ['Positions Table', 'All open positions across all accounts with P&L'],
            ['P&L Breakdown', 'By account and by symbol tables'],
            ['Auto-Refresh', 'Updates every 30 seconds'],
        ],
        col_widths=[2.0, 4.5]
    )

    doc.add_page_break()

    # ── 7. Deployment ──
    add_heading_styled(doc, '7. Deployment Options')

    add_styled_table(doc,
        ['Method', 'Command', 'Use Case'],
        [
            ['Local Python', 'python app.py --mode paper', 'Development and testing'],
            ['Local Live', 'python app.py --mode live', 'Production with MT5 on same machine'],
            ['Docker', 'docker-compose up -d', 'Containerized deployment'],
            ['Kubernetes', 'kubectl apply -f k8s/', 'Cloud-scale with auto-scaling (1-3 pods)'],
            ['Dashboard Only', 'python app.py --dashboard-only', 'Monitoring without trading'],
        ],
        col_widths=[1.5, 2.8, 2.2]
    )

    add_heading_styled(doc, '7.1 Docker Services', level=2)
    add_bullet(doc, 'reysentinel_engine: Main trading engine + dashboards')
    add_bullet(doc, 'reysentinel_heatmap: Portfolio heatmap dashboard')
    add_bullet(doc, 'reysentinel_multi: Multi-account dashboard')
    add_bullet(doc, 'redis: Signal caching for sub-second access')

    add_heading_styled(doc, '7.2 Kubernetes', level=2)
    add_bullet(doc, 'Deployment with liveness/readiness probes')
    add_bullet(doc, 'ClusterIP + NodePort services (30060-30062)')
    add_bullet(doc, 'ConfigMap for config.yaml')
    add_bullet(doc, 'HPA: auto-scale 1-3 replicas at 70% CPU')

    doc.add_page_break()

    # ── 8. Mobile App ──
    add_heading_styled(doc, '8. Mobile App')
    add_body(doc, 'React Native (Expo) app for iOS and Android:')

    add_styled_table(doc,
        ['Screen', 'Features'],
        [
            ['Dashboard', 'Portfolio summary cards, daily P&L chart, quick action buttons'],
            ['Positions', 'List of open positions with color-coded P&L, pull-to-refresh'],
            ['Alerts', 'Regime changes, correlation breaks, trade executions with severity colors'],
            ['Settings', 'API URL configuration, connection test, account selection'],
        ],
        col_widths=[1.5, 5.0]
    )

    add_body(doc, 'Components: TradeCard (position display), PnLChart (line chart), QuickActions (close-all, pause, override)')

    doc.add_page_break()

    # ── 9. How It All Works Together ──
    add_heading_styled(doc, '9. How It All Works Together')

    add_heading_styled(doc, 'Scenario: Trading Cycle on BTCUSD', level=2)
    add_body(doc, 'Here is a concrete example of one trading cycle:')

    add_body(doc, '1. The engine fetches 500 bars of BTCUSD 1-minute data from MT5.')
    add_body(doc, '2. The indicator engine computes RSI=28, MACD hist positive, Ichimoku bullish, volume 2.3x average.')
    add_body(doc, '3. The regime detector classifies the market as TRENDING_UP (ADX=32, strong directional bias).')
    add_body(doc, '4. The volume profile shows price near POC with VAL support below.')
    add_body(doc, '5. The sentiment bridge reports +0.45 bullish score (Reddit + CoinDesk positive).')
    add_body(doc, '6. Three models vote:')
    add_bullet(doc, 'Gemma 4: BUY at 0.85 confidence (RSI oversold + Ichimoku + volume)')
    add_bullet(doc, 'LLaMA 3: BUY at 0.72 confidence (conservative, wants more confirmation)')
    add_bullet(doc, 'ONNX: BUY at 0.68 confidence (indicator scoring)')
    add_body(doc, '7. Ensemble: Weighted average = BUY at 0.77 confidence. All 3 agree.')
    add_body(doc, '8. Risk manager approves: confidence > threshold, no open BTCUSD position, daily loss OK.')
    add_body(doc, '9. Position size: $1,000 risk (1% of $100K) / ATR-based SL = 0.02 lots.')
    add_body(doc, '10. MT5 executes: BUY 0.02 BTCUSD @ 67,450 | SL: 67,200 | TP: 67,825')
    add_body(doc, '11. Trade logged to journal, dashboard updated, outcome tracked for model calibration.')

    doc.add_page_break()

    # ── 10. Current Status ──
    add_heading_styled(doc, '10. Current Status')

    add_styled_table(doc,
        ['Module', 'Status', 'Notes'],
        [
            ['Ensemble Engine', 'BUILT', 'Gemma + LLaMA + ONNX with 3 voting methods'],
            ['Regime Detector', 'BUILT', 'Rule-based + HMM option'],
            ['Sentiment Bridge', 'BUILT', 'Reddit + RSS + Fear&Greed scrapers'],
            ['Volume Profile', 'BUILT', 'TPO, POC, VAH/VAL, HVN/LVN'],
            ['Correlation Engine', 'BUILT', 'Rolling corr, lag detection, cointegration'],
            ['Pattern Recognition', 'BUILT', 'CNN architecture + rule-based fallback'],
            ['Heatmap Dashboard', 'BUILT', 'FastAPI + D3.js with dark/light theme'],
            ['Audit Logger', 'BUILT', 'JSONL + SQLite + CSV export + reconciler'],
            ['Multi-Account', 'BUILT', 'N-account manager + consolidated P&L'],
            ['Docker/K8s', 'BUILT', 'Dockerfile, compose, K8s manifests, HPA'],
            ['Mobile App', 'BUILT', 'React Native (Expo) with 4 screens'],
        ],
        col_widths=[2.0, 0.8, 3.7]
    )

    doc.add_page_break()

    # ── 11. Roadmap ──
    add_heading_styled(doc, '11. Roadmap & What We Can Build Next')

    add_heading_styled(doc, 'Immediate (This Week)', level=2)
    add_bullet(doc, 'Train the CNN pattern recognition model on historical data')
    add_bullet(doc, 'Connect sentiment bridge to live Reddit/RSS feeds with API keys')
    add_bullet(doc, 'Run end-to-end paper trading test with full ensemble')
    add_bullet(doc, 'Validate Docker deployment on local machine')

    add_heading_styled(doc, 'Short-Term (2-4 Weeks)', level=2)
    add_bullet(doc, 'Backtest ensemble vs single-model (Gemma alone) on 3 months of data')
    add_bullet(doc, 'Add Telegram/Discord notifications for trade alerts')
    add_bullet(doc, 'Integrate correlation alerts into ensemble decision (decorrelation = reduce size)')
    add_bullet(doc, 'Build mobile app APK and test on Android device')
    add_bullet(doc, 'Add order flow / Level 2 data integration for crypto')

    add_heading_styled(doc, 'Medium-Term (1-3 Months)', level=2)
    add_bullet(doc, 'Deploy to AWS EC2 / DigitalOcean with Docker Compose')
    add_bullet(doc, 'Add forex and commodity pairs (XAUUSD, EURUSD, US100)')
    add_bullet(doc, 'Implement trailing stop optimization using volume profile levels')
    add_bullet(doc, 'Build web-based backtesting UI for ensemble strategies')
    add_bullet(doc, 'Add portfolio optimization (Markowitz / Black-Litterman)')
    add_bullet(doc, 'Implement copy-trading: mirror signals across multiple accounts')

    add_heading_styled(doc, 'Long-Term (3-6 Months)', level=2)
    add_bullet(doc, 'Reinforcement learning agent for dynamic strategy selection')
    add_bullet(doc, 'Alternative data: on-chain metrics, options flow, whale tracking')
    add_bullet(doc, 'SaaS dashboard: multi-tenant platform for other traders')
    add_bullet(doc, 'Regulatory compliance suite for fund management')
    add_bullet(doc, 'iOS App Store / Google Play Store release')

    add_heading_styled(doc, 'Possible New Products', level=2)
    add_styled_table(doc,
        ['Product Idea', 'Description', 'Effort'],
        [
            ['ReySignals', 'Signal-as-a-service: publish ensemble signals via API/Telegram', '2-3 weeks'],
            ['ReyCopy', 'Copy-trading platform: mirror trades across N accounts', '3-4 weeks'],
            ['ReyBacktest', 'Web-based backtesting with ensemble strategy builder', '4-6 weeks'],
            ['ReyFund', 'Portfolio management suite with compliance for fund managers', '2-3 months'],
            ['ReyAlpha', 'Alpha research platform: test signals, rank strategies, factor analysis', '2-3 months'],
        ],
        col_widths=[1.5, 3.5, 1.5]
    )

    # Save
    path = os.path.join(os.path.dirname(__file__), 'ReySentinel_Architecture.docx')
    doc.save(path)
    print(f'Created: {path}')
    return path


if __name__ == '__main__':
    create_rey_capital_ai_bot_doc()
    create_reysentinel_doc()
    print('\nDone! Both documents created.')
